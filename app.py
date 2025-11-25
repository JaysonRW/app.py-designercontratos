import os
import tempfile
import re
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})

UPLOAD_FOLDER = tempfile.gettempdir()

def hex_to_rgb(hex_color):
    """Converte cor hexadecimal para RGB"""
    color = hex_color.lstrip('#')
    return tuple(int(color[i:i+2], 16) for i in (0, 2, 4))

def shade_cell(cell, hex_color):
    """Aplica cor de fundo em célula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), hex_color.lstrip('#'))
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_formatted_document(text_content, primary_color, logo_path=None):
    """
    Cria documento Word formatado DO ZERO a partir de texto
    """
    doc = Document()
    
    # Configurações de página
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    
    # Cores
    r_primary, g_primary, b_primary = hex_to_rgb(primary_color)
    
    # 1. CABEÇALHO COM LOGO
    if logo_path and os.path.exists(logo_path):
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(1.8))
        
        # Linha separadora
        separator = header.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = separator.add_run('_' * 80)
        sep_run.font.color.rgb = RGBColor(r_primary, g_primary, b_primary)
        sep_run.font.size = Pt(6)
    
    # 2. PROCESSA O TEXTO
    lines = text_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Linha vazia - adiciona espaçamento
            doc.add_paragraph()
            continue
        
        # Detecta TÍTULO (texto em MAIÚSCULAS ou curto e centralizado)
        is_title = (
            line.isupper() and len(line) < 100 or
            line.startswith('CONTRATO') or
            line.startswith('CLÁUSULA') or
            re.match(r'^\d+\.', line)  # Começa com número (ex: "1. Título")
        )
        
        # Detecta item de lista
        is_list_item = line.startswith(('- ', '• ', '* '))
        
        # Detecta tabela simples (formato: CHAVE: Valor)
        is_table_row = ':' in line and len(line.split(':')) == 2
        
        if is_title:
            # TÍTULOS
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line)
            run.font.color.rgb = RGBColor(r_primary, g_primary, b_primary)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = 'Arial'
            
        elif is_list_item:
            # ITENS DE LISTA
            para = doc.add_paragraph(line, style='List Bullet')
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(50, 50, 50)
                
        else:
            # TEXTO NORMAL
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(line)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(50, 50, 50)
            run.font.name = 'Arial'
    
    # 3. RODAPÉ
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Documento gerado automaticamente | Designer de Contratos')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    
    return doc

def create_table_document(text_content, primary_color, logo_path=None):
    """
    Versão com detecção automática de tabelas
    """
    doc = Document()
    
    # Configurações de página
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    
    # Cores
    r_primary, g_primary, b_primary = hex_to_rgb(primary_color)
    
    # Cabeçalho com logo
    if logo_path and os.path.exists(logo_path):
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(1.8))
    
    # Processa linhas
    lines = text_content.split('\n')
    table_data = []
    in_table = False
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Se estava em uma tabela, cria ela
            if in_table and table_data:
                create_formatted_table(doc, table_data, primary_color)
                table_data = []
                in_table = False
            doc.add_paragraph()
            continue
        
        # Detecta linha de tabela (formato: CHAVE: Valor)
        if ':' in line:
            parts = line.split(':', 1)
            if len(parts) == 2:
                table_data.append([parts[0].strip(), parts[1].strip()])
                in_table = True
                continue
        
        # Se estava em tabela mas linha não é tabela, cria a tabela
        if in_table and table_data:
            create_formatted_table(doc, table_data, primary_color)
            table_data = []
            in_table = False
        
        # Adiciona parágrafo normal
        is_title = line.isupper() or line.startswith('CONTRATO')
        
        para = doc.add_paragraph()
        if is_title:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        run = para.add_run(line)
        run.font.size = Pt(14 if is_title else 11)
        run.font.bold = is_title
        
        if is_title:
            run.font.color.rgb = RGBColor(r_primary, g_primary, b_primary)
        else:
            run.font.color.rgb = RGBColor(50, 50, 50)
    
    # Cria última tabela se houver
    if table_data:
        create_formatted_table(doc, table_data, primary_color)
    
    return doc

def create_formatted_table(doc, data, primary_color):
    """Cria tabela formatada"""
    r, g, b = hex_to_rgb(primary_color)
    
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Light Grid Accent 1'
    
    for i, (key, value) in enumerate(data):
        row = table.rows[i]
        
        # Célula da chave
        key_cell = row.cells[0]
        key_para = key_cell.paragraphs[0]
        key_run = key_para.add_run(key)
        key_run.font.bold = True
        key_run.font.size = Pt(10)
        shade_cell(key_cell, primary_color)
        key_run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Célula do valor
        value_cell = row.cells[1]
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(value)
        value_run.font.size = Pt(10)
        
        # Zebrado
        if i % 2 == 1:
            shade_cell(value_cell, '#F3F4F6')

@app.route('/api/generate', methods=['POST'])
def generate_document():
    """
    Gera documento Word formatado a partir de texto
    """
    try:
        data = request.get_json()
        
        text_content = data.get('text', '')
        primary_color = data.get('primaryColor', '#4F46E5')
        
        if not text_content:
            return jsonify({'error': 'Texto não fornecido'}), 400
        
        # Processa logo se enviado (como base64)
        logo_path = None
        if 'logoBase64' in data:
            import base64
            logo_data = base64.b64decode(data['logoBase64'].split(',')[1])
            logo_path = os.path.join(UPLOAD_FOLDER, 'logo_temp.png')
            with open(logo_path, 'wb') as f:
                f.write(logo_data)
        
        # Gera documento
        doc = create_table_document(text_content, primary_color, logo_path)
        
        # Salva
        output_filename = 'contrato_formatado.docx'
        output_path = os.path.join(UPLOAD_FOLDER, output_filename)
        doc.save(output_path)
        
        # Limpa logo temporário
        if logo_path and os.path.exists(logo_path):
            os.remove(logo_path)
        
        base_url = request.url_root.rstrip('/')
        
        return jsonify({
            'docxUrl': f"{base_url}/api/download?file={output_filename}",
            'message': 'Documento gerado com sucesso!'
        })
        
    except Exception as e:
        print(f"Erro: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/download', methods=['GET'])
def download_file():
    filename = request.args.get('file')
    if not filename:
        return "Arquivo não especificado", 400
    
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    
    try:
        response = send_file(file_path, as_attachment=True)
        
        @response.call_on_close
        def cleanup():
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except:
                pass
        
        return response
    except FileNotFoundError:
        return "Arquivo não encontrado", 404

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
